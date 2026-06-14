import streamlit as st
import pandas as pd
import numpy as np
import re
import os
import json
import time
import random
import hashlib
import tempfile
from datetime import datetime
from pathlib import Path

from groq import Groq

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import SystemMessage, HumanMessage
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader

from docx import Document as DocxDocument

# =========================
# PAGE CONFIG
# =========================

st.set_page_config(
    page_title="Sales Field Intelligence Chatbot",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# =========================
# CONFIG: COLUMN ALIASES
# =========================

COLUMN_ALIASES = {
    "sales": [
        "sales", "nama sales", "petugas", "marketing", "nama marketing"
    ],
    "supervisor": [
        "supervisor", "leader", "nama supervisor", "pic supervisor", "koordinator"
    ],
    "region": [
        "region", "regional", "ro", "region office", "regional office"
    ],
    "area": [
        "area", "nama area", "wilayah"
    ],
    "branch": [
        "branch", "cabang", "kc", "kcp", "unit", "nama cabang", "outlet"
    ],
    "stage": [
        "stage", "stage pipeline", "pipeline stage", "status pipeline",
        "pipeline", "status", "kategori pipeline", "status nasabah",
        "tahapan", "tahap pipeline", "tahapan pipeline", "status follow up", "Stage Pipeline"
    ],
    "activity": [
        "activity", "aktivitas", "jumlah activity", "jumlah aktivitas",
        "kunjungan", "jumlah kunjungan", "follow up", "followup",
        "jumlah follow up", "jml activity"
    ],
    "result": [
        "status cair", "hasil", "status hasil", "status akhir",
        "cair", "status pencairan", "hasil cair"
    ],
    "plafond": [
        "plafond", "nominal", "nilai", "potensi", "nilai potensi",
        "amount", "pengajuan", "nominal pengajuan"
    ],
    "period": [
        "periode", "tanggal", "tgl", "date", "bulan", "tanggal update",
        "tgl update"
    ],
    "constraint": [
        "kendala", "keterangan kendala", "alasan", "reason",
        "alasan batal", "alasan reject", "alasan batal reject",
        "alasan batal/reject", "keterangan"
    ]
}

# =========================
# CONFIG: STAGE ALIASES
# =========================

STAGE_ALIASES = {
    "tidak bertemu": "Tidak Bertemu",
    "tbn": "Tidak Bertemu",
    "belum bertemu": "Tidak Bertemu",
    "not meet": "Tidak Bertemu",

    "dingin": "Dingin",
    "cool": "Dingin",
    "cold": "Dingin",

    "hangat": "Hangat",
    "warm": "Hangat",

    "panas": "Panas",
    "hot": "Panas",

    "pemberkasan": "Pemberkasan",
    "berkas": "Pemberkasan",
    "incoming": "Pemberkasan",
    "proses": "Pemberkasan",
    "proses berkas": "Pemberkasan",

    "perjanjian": "Perjanjian",
    "akad": "Perjanjian",
    "agreement": "Perjanjian",

    "cair": "Cair",
    "realized": "Cair",
    "disburse": "Cair",
    "disbursed": "Cair"
}

STAGE_ORDER = {
    "Tidak Bertemu": 1,
    "Dingin": 2,
    "Hangat": 3,
    "Panas": 4,
    "Pemberkasan": 5,
    "Perjanjian": 6,
    "Cair": 7
}

NEGATIVE_CAIR_KEYWORDS = [
    "belum cair",
    "batal cair",
    "tidak cair",
    "gagal cair",
    "reject",
    "ditolak",
    "cancel",
    "canceled",
    "cancelled"
]

POSITIVE_CAIR_PATTERNS = [
    r"^cair$",
    r"\bsudah cair\b",
    r"\btelah cair\b",
    r"\bterealisasi\b",
    r"\brealized\b",
    r"\bdisbursed\b"
]

# =========================
# SESSION STATE INIT
# =========================

def init_session_state():
    defaults = {
        "messages": [],
        "excel_df": None,
        "excel_ready": False,
        "excel_processed": False,
        "detected_columns": {},
        "df_pipeline": None,
        "overall_summary": None,
        "summary_sales": None,
        "summary_supervisor": None,
        "summary_region": None,
        "summary_branch": None,
        "coaching_sales": None,
        "coaching_supervisor": None,
        "pending_prompt": None,
        "chat_export_json": None,
        "chat_export_markdown": None,
        "rule_insights": [],
        "excel_analysis_context": None,
        "rag_ready": False,
        "rag_vector_store": None,
        "rag_sources": [],
        "excel_answer_cache": {},
        "router_cache": {},
        "last_route_info": None,
        "api_status": {
            "gemini": False,
            "groq": False
        }
    }

    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


init_session_state()

# =========================
# EXCEL ANALYZER FUNCTIONS
# =========================

def normalize_col_name(col):
    col = str(col).strip().lower()
    col = re.sub(r"\s+", " ", col)
    col = col.replace("_", " ")
    return col


def detect_columns(df, column_aliases=COLUMN_ALIASES):
    normalized_columns = {
        normalize_col_name(col): col
        for col in df.columns
    }

    detected = {}

    for standard_name, aliases in column_aliases.items():
        detected[standard_name] = None

        for alias in aliases:
            alias_norm = normalize_col_name(alias)

            if alias_norm in normalized_columns:
                detected[standard_name] = normalized_columns[alias_norm]
                break

    return detected


def clean_text_value(value):
    if pd.isna(value):
        return ""
    value = str(value).strip()
    value = re.sub(r"\s+", " ", value)
    return value


def normalize_stage_value(value):
    raw = clean_text_value(value)
    raw_lower = raw.lower()

    if not raw_lower:
        return "Tidak Diketahui"

    if raw_lower in STAGE_ALIASES:
        return STAGE_ALIASES[raw_lower]

    for key, canonical in STAGE_ALIASES.items():
        if key in raw_lower:
            return canonical

    return raw.title()


def to_number(value):
    if pd.isna(value):
        return 0

    if isinstance(value, (int, float, np.integer, np.floating)):
        return float(value)

    value = str(value)
    value = value.replace("Rp", "").replace("rp", "")
    value = value.replace(" ", "")
    value = value.replace(".", "")
    value = value.replace(",", ".")

    try:
        return float(value)
    except Exception:
        return 0


def contains_negative_cair_text(value):
    text = clean_text_value(value).lower()
    return any(keyword in text for keyword in NEGATIVE_CAIR_KEYWORDS)


def is_positive_cair_text(value):
    text = clean_text_value(value).lower()

    if not text:
        return False

    if contains_negative_cair_text(text):
        return False

    return any(re.search(pattern, text) for pattern in POSITIVE_CAIR_PATTERNS)


def prepare_pipeline_dataframe(df, detected):
    data = pd.DataFrame()

    def get_col(key):
        return detected.get(key)

    for key in ["sales", "supervisor", "region", "area", "branch", "period", "constraint"]:
        col = get_col(key)

        if col and col in df.columns:
            data[key] = df[col].apply(clean_text_value)
        else:
            data[key] = ""

    stage_col = get_col("stage")

    if stage_col and stage_col in df.columns:
        data["stage_raw"] = df[stage_col].apply(clean_text_value)
        data["stage"] = df[stage_col].apply(normalize_stage_value)
    else:
        data["stage_raw"] = ""
        data["stage"] = "Tidak Diketahui"

    activity_col = get_col("activity")

    if activity_col and activity_col in df.columns:
        data["activity"] = df[activity_col].apply(to_number)
    else:
        data["activity"] = 1

    result_col = get_col("result")

    if result_col and result_col in df.columns:
        data["result"] = df[result_col].apply(clean_text_value)
    else:
        data["result"] = ""

    plafond_col = get_col("plafond")

    if plafond_col and plafond_col in df.columns:
        data["plafond"] = df[plafond_col].apply(to_number)
    else:
        data["plafond"] = 0

    data["stage_score"] = data["stage"].map(STAGE_ORDER).fillna(0)

    data["is_cair"] = data.apply(
        lambda row: (
            str(row["stage"]).lower() == "cair"
            or is_positive_cair_text(row.get("result", ""))
        ),
        axis=1
    )

    data["is_early_stage"] = data["stage"].isin(["Tidak Bertemu", "Dingin"])
    data["is_middle_stage"] = data["stage"].isin(["Hangat"])
    data["is_mature_stage"] = data["stage"].isin(["Panas", "Pemberkasan", "Perjanjian"])
    data["is_final_stage"] = data["is_cair"]

    return data


def fix_cair_flags(df):
    df_fixed = df.copy()

    negative_stage_mask = (
        df_fixed["stage"].str.lower().eq("cair")
        &
        df_fixed["stage_raw"].apply(contains_negative_cair_text)
    )

    df_fixed.loc[negative_stage_mask, "stage"] = "Tidak Diketahui"

    df_fixed["stage_score"] = df_fixed["stage"].map(STAGE_ORDER).fillna(0)

    df_fixed["is_cair"] = df_fixed.apply(
        lambda row: (
            str(row["stage"]).lower() == "cair"
            or is_positive_cair_text(row.get("result", ""))
        ),
        axis=1
    )

    df_fixed["is_early_stage"] = df_fixed["stage"].isin(["Tidak Bertemu", "Dingin"])
    df_fixed["is_middle_stage"] = df_fixed["stage"].isin(["Hangat"])
    df_fixed["is_mature_stage"] = df_fixed["stage"].isin(["Panas", "Pemberkasan", "Perjanjian"])
    df_fixed["is_final_stage"] = df_fixed["is_cair"]

    return df_fixed


def summarize_overall_pipeline(df):
    total_rows = len(df)
    total_activity = df["activity"].sum()
    total_plafond = df["plafond"].sum()
    total_cair = df["is_cair"].sum()
    total_cair_plafond = df.loc[df["is_cair"], "plafond"].sum()

    stage_summary = (
        df.groupby("stage", dropna=False)
        .agg(
            jumlah_pipeline=("stage", "size"),
            total_activity=("activity", "sum"),
            total_plafond=("plafond", "sum"),
            jumlah_cair=("is_cair", "sum")
        )
        .reset_index()
    )

    stage_summary["stage_order"] = stage_summary["stage"].map(STAGE_ORDER).fillna(99)
    stage_summary = stage_summary.sort_values("stage_order").drop(columns=["stage_order"])

    early_count = df["is_early_stage"].sum()
    middle_count = df["is_middle_stage"].sum()
    mature_count = df["is_mature_stage"].sum()
    final_count = df["is_final_stage"].sum()

    avg_stage_score = df["stage_score"].replace(0, np.nan).mean()

    result = {
        "total_rows": int(total_rows),
        "total_activity": float(total_activity),
        "total_plafond": float(total_plafond),
        "total_cair": int(total_cair),
        "total_cair_plafond": float(total_cair_plafond),
        "early_count": int(early_count),
        "middle_count": int(middle_count),
        "mature_count": int(mature_count),
        "final_count": int(final_count),
        "avg_stage_score": float(avg_stage_score) if not pd.isna(avg_stage_score) else 0,
        "stage_summary": stage_summary
    }

    return result


def summarize_by_dimension(df, dimension):
    if dimension not in df.columns:
        return pd.DataFrame()

    if df[dimension].replace("", np.nan).dropna().empty:
        return pd.DataFrame()

    grouped = (
        df.groupby(dimension, dropna=False)
        .agg(
            jumlah_pipeline=("stage", "size"),
            total_activity=("activity", "sum"),
            total_plafond=("plafond", "sum"),
            jumlah_cair=("is_cair", "sum"),
            avg_stage_score=("stage_score", "mean"),
            early_stage=("is_early_stage", "sum"),
            middle_stage=("is_middle_stage", "sum"),
            mature_stage=("is_mature_stage", "sum"),
            final_stage=("is_final_stage", "sum")
        )
        .reset_index()
    )

    grouped["plafond_cair"] = grouped[dimension].apply(
        lambda val: df.loc[(df[dimension] == val) & (df["is_cair"]), "plafond"].sum()
    )

    grouped["activity_per_pipeline"] = (
        grouped["total_activity"] / grouped["jumlah_pipeline"].replace(0, np.nan)
    )

    grouped["cair_ratio"] = (
        grouped["jumlah_cair"] / grouped["jumlah_pipeline"].replace(0, np.nan)
    )

    grouped = grouped.fillna(0)

    return grouped.sort_values(
        by=["jumlah_pipeline", "total_activity"],
        ascending=[False, False]
    )


def generate_rule_based_insight(overall):
    insights = []

    total = overall["total_rows"]

    if total == 0:
        return ["Data kosong, belum dapat dianalisis."]

    early_pct = overall["early_count"] / total
    mature_pct = overall["mature_count"] / total
    cair_pct = overall["total_cair"] / total

    stage_summary = overall["stage_summary"]

    if not stage_summary.empty:
        dominant_stage_row = stage_summary.sort_values("jumlah_pipeline", ascending=False).iloc[0]
        dominant_stage = dominant_stage_row["stage"]
        dominant_count = dominant_stage_row["jumlah_pipeline"]

        insights.append(
            f"Stage dominan adalah {dominant_stage} dengan {int(dominant_count):,} pipeline."
        )

    if early_pct >= 0.5:
        insights.append(
            "Pipeline masih didominasi tahap awal Tidak Bertemu/Dingin, sehingga fokus utama perlu diarahkan pada validasi database dan penguatan follow-up awal."
        )

    if mature_pct >= 0.4 and cair_pct < 0.2:
        insights.append(
            "Pipeline matang relatif tersedia, tetapi rasio Cair masih rendah. Hal ini mengindikasikan perlunya validasi bottleneck pada tahap Panas, Pemberkasan, atau Perjanjian."
        )

    if overall["total_activity"] > 0 and overall["total_cair"] == 0:
        insights.append(
            "Activity sudah tercatat, tetapi belum terdapat pencairan. Efektivitas activity perlu dilihat dari pergerakan stage pipeline."
        )

    if overall["total_cair"] > 0:
        insights.append(
            f"Terdapat {overall['total_cair']:,} pipeline berstatus Cair. Analisis lanjutan perlu melihat apakah pipeline baru tetap tersedia untuk menjaga kesinambungan hasil."
        )

    if not insights:
        insights.append(
            "Data sudah terbaca, namun belum menunjukkan pola dominan yang kuat. Perlu analisis berdasarkan Sales, Supervisor, atau Region."
        )

    return insights


def build_coaching_priority(summary_df, name_col, top_n=10):
    if summary_df is None or summary_df.empty:
        return pd.DataFrame()

    df = summary_df.copy()

    median_activity = df["total_activity"].median()

    df["mature_not_cair"] = df["mature_stage"]

    df["activity_high_no_cair"] = (
        (df["total_activity"] >= median_activity)
        &
        (df["jumlah_cair"] == 0)
    ).astype(int)

    df["activity_low"] = (
        df["total_activity"] < median_activity
    ).astype(int)

    df["coaching_score"] = (
        df["early_stage"] * 1
        + df["mature_not_cair"].clip(lower=0) * 2
        + df["activity_high_no_cair"] * 3
        + df["activity_low"] * 1
    )

    df["indikasi"] = ""

    df.loc[df["early_stage"] > df["mature_stage"], "indikasi"] += "Pipeline masih banyak di tahap awal. "
    df.loc[df["mature_not_cair"] > 0, "indikasi"] += "Terdapat pipeline matang yang belum cair. "
    df.loc[df["activity_high_no_cair"] == 1, "indikasi"] += "Activity relatif tinggi tetapi belum menghasilkan Cair. "
    df.loc[df["activity_low"] == 1, "indikasi"] += "Activity relatif rendah dibanding kelompok. "

    result = df.sort_values("coaching_score", ascending=False).head(top_n)

    cols = [
        name_col,
        "jumlah_pipeline",
        "total_activity",
        "jumlah_cair",
        "early_stage",
        "middle_stage",
        "mature_stage",
        "final_stage",
        "coaching_score",
        "indikasi"
    ]

    available_cols = [col for col in cols if col in result.columns]

    return result[available_cols]


def dataframe_to_markdown_limited(df, max_rows=10):
    if df is None or df.empty:
        return "Tidak tersedia."

    limited_df = df.head(max_rows)

    try:
        return limited_df.to_markdown(index=False)
    except Exception:
        return limited_df.to_string(index=False)


def build_excel_analysis_context(
    overall,
    summary_sales=None,
    summary_supervisor=None,
    summary_region=None,
    coaching_sales=None,
    coaching_supervisor=None,
    max_rows=10
):
    stage_summary_md = dataframe_to_markdown_limited(overall["stage_summary"], max_rows=20)
    sales_md = dataframe_to_markdown_limited(summary_sales, max_rows=max_rows)
    supervisor_md = dataframe_to_markdown_limited(summary_supervisor, max_rows=max_rows)
    region_md = dataframe_to_markdown_limited(summary_region, max_rows=max_rows)
    coaching_sales_md = dataframe_to_markdown_limited(coaching_sales, max_rows=max_rows)
    coaching_supervisor_md = dataframe_to_markdown_limited(coaching_supervisor, max_rows=max_rows)

    context = f"""
## Overall Summary

Total baris data: {overall['total_rows']}
Total activity: {overall['total_activity']}
Total plafond: {overall['total_plafond']}
Jumlah Cair: {overall['total_cair']}
Total plafond Cair: {overall['total_cair_plafond']}
Pipeline tahap awal: {overall['early_count']}
Pipeline tahap tengah: {overall['middle_count']}
Pipeline tahap matang: {overall['mature_count']}
Pipeline tahap final/Cair: {overall['final_count']}
Rata-rata stage score: {overall['avg_stage_score']}

## Stage Summary

{stage_summary_md}

## Summary by Region

{region_md}

## Summary by Supervisor

{supervisor_md}

## Summary by Sales

{sales_md}

## Prioritas Coaching Sales

{coaching_sales_md}

## Prioritas Coaching Supervisor

{coaching_supervisor_md}
"""

    return context.strip()


def run_excel_analyzer(df, detected_columns):
    df_pipeline = prepare_pipeline_dataframe(df, detected_columns)
    df_pipeline = fix_cair_flags(df_pipeline)

    overall_summary = summarize_overall_pipeline(df_pipeline)

    summary_sales = summarize_by_dimension(df_pipeline, "sales")
    summary_supervisor = summarize_by_dimension(df_pipeline, "supervisor")
    summary_region = summarize_by_dimension(df_pipeline, "region")
    summary_branch = summarize_by_dimension(df_pipeline, "branch")

    rule_insights = generate_rule_based_insight(overall_summary)

    coaching_sales = build_coaching_priority(summary_sales, "sales", top_n=10)
    coaching_supervisor = build_coaching_priority(summary_supervisor, "supervisor", top_n=10)

    excel_context = build_excel_analysis_context(
        overall=overall_summary,
        summary_sales=summary_sales,
        summary_supervisor=summary_supervisor,
        summary_region=summary_region,
        coaching_sales=coaching_sales,
        coaching_supervisor=coaching_supervisor,
        max_rows=10
    )

    return {
        "df_pipeline": df_pipeline,
        "overall_summary": overall_summary,
        "summary_sales": summary_sales,
        "summary_supervisor": summary_supervisor,
        "summary_region": summary_region,
        "summary_branch": summary_branch,
        "rule_insights": rule_insights,
        "coaching_sales": coaching_sales,
        "coaching_supervisor": coaching_supervisor,
        "excel_analysis_context": excel_context
    }

# =========================
# AI / RAG / ROUTER FUNCTIONS
# =========================

GROQ_MODEL_LIGHT = "llama-3.1-8b-instant"
GROQ_MODEL_MEDIUM = "llama-3.3-70b-versatile"
GEMINI_MODEL_HEAVY = "gemini-2.5-flash"
EMBEDDING_MODEL_NAME = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"


def make_cache_key(payload):
    raw = json.dumps(
        payload,
        ensure_ascii=False,
        sort_keys=True,
        default=str
    )
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def is_rate_limit_error(error):
    err_text = str(error).lower()

    keywords = [
        "429",
        "resource_exhausted",
        "rate limit",
        "quota",
        "too many requests"
    ]

    return any(keyword in err_text for keyword in keywords)


def run_with_retry(func, max_retries=3, base_delay=2):
    for attempt in range(max_retries):
        try:
            return func()

        except Exception as e:
            if is_rate_limit_error(e) and attempt < max_retries - 1:
                wait_time = base_delay * (2 ** attempt) + random.uniform(0, 1)
                time.sleep(wait_time)
            else:
                raise e


@st.cache_resource(show_spinner=False)
def get_embedding_model():
    return HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL_NAME
    )


def build_tone_instruction(tone_mode):
    if tone_mode == "Ringkas Eksekutif":
        return """
Gunakan gaya ringkas eksekutif.
Jawaban harus padat, formal, dan fokus pada insight utama, risiko, serta rekomendasi prioritas.
"""

    if tone_mode == "Detail Analitis":
        return """
Gunakan gaya detail analitis.
Jawaban boleh lebih panjang, terstruktur, dan menjelaskan dasar data, interpretasi, risiko, serta tindak lanjut.
"""

    return """
Gunakan gaya formal operasional.
Jawaban harus jelas, profesional, mudah dipahami, dan berorientasi tindak lanjut.
"""


def call_groq_light_streamlit(question, groq_api_key, tone_mode="Formal Operasional"):
    if not groq_api_key:
        raise ValueError("Groq API Key belum tersedia.")

    client = Groq(api_key=groq_api_key)

    system_prompt = f"""
Anda adalah asisten ringan untuk Sales Field Intelligence Chatbot.

Tugas Anda:
- Menjawab sapaan.
- Membantu parafrase.
- Membuat ringkasan pendek.
- Membuat judul/caption.
- Menjelaskan hal sederhana.

Batasan:
- Jangan melakukan analisis data berat.
- Jangan mengarang angka, nama, atau kesimpulan.
- Jika pertanyaan membutuhkan analisis Excel, RAG, pipeline, activity, atau coaching berbasis data, arahkan ke mode analisis.

{build_tone_instruction(tone_mode)}
"""

    def _call():
        completion = client.chat.completions.create(
            model=GROQ_MODEL_LIGHT,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": question}
            ],
            temperature=0.3,
            max_tokens=700
        )

        return completion.choices[0].message.content

    return run_with_retry(_call, max_retries=2, base_delay=1)


def call_gemini_direct_streamlit(question, gemini_api_key, system_prompt=None, tone_mode="Formal Operasional"):
    if not gemini_api_key:
        raise ValueError("Gemini API Key belum tersedia.")

    if system_prompt is None:
        system_prompt = f"""
Anda adalah Sales Field Intelligence Chatbot.

Aturan:
1. Gunakan bahasa Indonesia formal dan jelas.
2. Jangan mengarang data.
3. Jika data atau konteks tidak tersedia, sampaikan keterbatasannya.
4. Jangan menyimpulkan penyebab pasti jika data pendukung tidak tersedia.
5. Jawaban harus berorientasi insight dan tindak lanjut.

{build_tone_instruction(tone_mode)}
"""

    llm = ChatGoogleGenerativeAI(
        google_api_key=gemini_api_key,
        model=GEMINI_MODEL_HEAVY,
        temperature=0.2,
        max_output_tokens=1200
    )

    def _call():
        response = llm.invoke([
            SystemMessage(content=system_prompt),
            HumanMessage(content=question)
        ])

        return response.content

    return run_with_retry(_call, max_retries=3, base_delay=2)


def ask_excel_analyzer_streamlit(question, gemini_api_key, tone_mode="Formal Operasional", use_cache=True):
    if not st.session_state.excel_processed:
        return {
            "answer": "Excel Analyzer belum aktif. Upload Excel dan klik Proses Excel Analyzer terlebih dahulu.",
            "sources": [],
            "from_cache": False
        }

    context = st.session_state.excel_analysis_context

    cache_key = make_cache_key({
        "type": "excel_analyzer",
        "question": question,
        "context_hash": hashlib.md5(context.encode("utf-8")).hexdigest(),
        "tone_mode": tone_mode
    })

    if use_cache and cache_key in st.session_state.excel_answer_cache:
        return {
            "answer": st.session_state.excel_answer_cache[cache_key],
            "sources": ["Excel summary dataframe"],
            "from_cache": True
        }

    system_prompt = f"""
Anda adalah Sales Field Intelligence Chatbot yang membantu menganalisis data pipeline dan activity.

Aturan wajib:
1. Jawab hanya berdasarkan summary data Excel yang diberikan.
2. Jangan mengarang nama Sales, Supervisor, Region, angka, atau nominal.
3. Jika data tidak tersedia, sampaikan keterbatasannya.
4. Jangan menyimpulkan penyebab pasti jika data tidak memuat kolom kendala.
5. Untuk rekomendasi coaching, gunakan struktur:
   - Ringkasan Temuan
   - Dasar Data
   - Indikasi Masalah
   - Rekomendasi Tindak Lanjut
   - Catatan Keterbatasan

{build_tone_instruction(tone_mode)}
"""

    prompt = f"""
Berikut adalah summary data Excel pipeline dan activity:

{context}

Pertanyaan user:
{question}
"""

    answer = call_gemini_direct_streamlit(
        question=prompt,
        gemini_api_key=gemini_api_key,
        system_prompt=system_prompt,
        tone_mode=tone_mode
    )

    if use_cache:
        st.session_state.excel_answer_cache[cache_key] = answer

    return {
        "answer": answer,
        "sources": ["Excel summary dataframe"],
        "from_cache": False
    }


def load_uploaded_knowledge_files(uploaded_files):
    docs = []

    if not uploaded_files:
        return docs

    for uploaded_file in uploaded_files:
        suffix = Path(uploaded_file.name).suffix.lower()
        file_bytes = uploaded_file.getvalue()

        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            if suffix == ".pdf":
                loader = PyPDFLoader(tmp_path)
                file_docs = loader.load()

                for doc in file_docs:
                    doc.metadata["source"] = uploaded_file.name
                    doc.metadata["type"] = "pdf"

                docs.extend(file_docs)

            elif suffix in [".txt", ".md"]:
                text = file_bytes.decode("utf-8", errors="ignore")

                docs.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": uploaded_file.name,
                            "type": suffix.replace(".", "")
                        }
                    )
                )

            elif suffix == ".docx":
                docx = DocxDocument(tmp_path)
                text_parts = []

                for paragraph in docx.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        text_parts.append(text)

                docs.append(
                    Document(
                        page_content="\n".join(text_parts),
                        metadata={
                            "source": uploaded_file.name,
                            "type": "docx"
                        }
                    )
                )

        finally:
            try:
                os.remove(tmp_path)
            except Exception:
                pass

    return docs


def build_rag_vector_store_streamlit(uploaded_files):
    docs = load_uploaded_knowledge_files(uploaded_files)

    if not docs:
        raise ValueError("Tidak ada dokumen knowledge base yang berhasil dibaca.")

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=900,
        chunk_overlap=120,
        separators=["\n\n", "\n", ".", " ", ""]
    )

    chunks = text_splitter.split_documents(docs)

    if not chunks:
        raise ValueError("Dokumen knowledge base tidak menghasilkan chunk.")

    embedding_model = get_embedding_model()

    vector_store = FAISS.from_documents(
        documents=chunks,
        embedding=embedding_model
    )

    sources = []

    for doc in docs:
        source = doc.metadata.get("source", "unknown")
        if source not in sources:
            sources.append(source)

    return vector_store, sources, len(chunks)


def format_docs_with_source(docs):
    context_parts = []

    for i, doc in enumerate(docs, start=1):
        source = doc.metadata.get("source", "unknown source")
        page = doc.metadata.get("page", None)
        doc_type = doc.metadata.get("type", "unknown")

        source_info = f"[Sumber {i}: {source}"

        if page is not None:
            source_info += f", halaman {page + 1}"

        source_info += f", tipe {doc_type}]"

        context_parts.append(
            f"{source_info}\n{doc.page_content}"
        )

    return "\n\n".join(context_parts)


def get_source_summary(docs):
    sources = []

    for doc in docs:
        source = doc.metadata.get("source", "unknown source")
        page = doc.metadata.get("page", None)

        if page is not None:
            label = f"{source} - halaman {page + 1}"
        else:
            label = source

        sources.append(label)

    return list(dict.fromkeys(sources))


def ask_rag_streamlit(question, gemini_api_key, tone_mode="Formal Operasional", k=4):
    if not st.session_state.rag_ready or st.session_state.rag_vector_store is None:
        return {
            "answer": (
                "Knowledge base RAG belum aktif. Upload dokumen knowledge base, lalu klik "
                "**Build RAG Knowledge Base** terlebih dahulu."
            ),
            "sources": [],
            "from_cache": False
        }

    vector_store = st.session_state.rag_vector_store

    docs = vector_store.similarity_search(question, k=k)

    context = format_docs_with_source(docs)

    system_prompt = f"""
Anda adalah Sales Field Intelligence Chatbot untuk membantu analisis pipeline, activity, dan rekomendasi coaching.

Aturan wajib:
1. Jawab hanya berdasarkan konteks knowledge base yang diberikan.
2. Jika konteks tidak cukup, katakan bahwa informasi belum tersedia di knowledge base.
3. Jangan mengarang penyebab, angka, nama, wilayah, atau aturan.
4. Jika ada istilah teknis, jelaskan secara sederhana.
5. Struktur jawaban:
   - Ringkasan Jawaban
   - Dasar Referensi
   - Rekomendasi
   - Catatan Keterbatasan

{build_tone_instruction(tone_mode)}
"""

    prompt = f"""
Konteks dari knowledge base:
{context}

Pertanyaan user:
{question}
"""

    answer = call_gemini_direct_streamlit(
        question=prompt,
        gemini_api_key=gemini_api_key,
        system_prompt=system_prompt,
        tone_mode=tone_mode
    )

    return {
        "answer": answer,
        "sources": get_source_summary(docs),
        "from_cache": False
    }


def should_use_excel_analyzer(question):
    q = question.lower().strip()

    strong_excel_keywords = [
        "data excel",
        "file excel",
        "berdasarkan data",
        "dari data",
        "pada data",
        "dataset",
        "sheet",
        "summary data",
        "ringkasan data",
        "data yang diupload",
        "data upload"
    ]

    data_analysis_keywords = [
        "analisis",
        "ringkasan",
        "summary",
        "berapa",
        "jumlah",
        "total",
        "ranking",
        "rank",
        "tertinggi",
        "terendah",
        "terbesar",
        "terkecil",
        "prioritas",
        "siapa",
        "mana",
        "bandingkan",
        "performa",
        "kinerja",
        "produktif",
        "tidak produktif"
    ]

    domain_keywords = [
        "pipeline",
        "activity",
        "aktivitas",
        "sales",
        "supervisor",
        "region",
        "cabang",
        "area",
        "coaching",
        "cair",
        "plafond",
        "panas",
        "pemberkasan",
        "perjanjian",
        "tidak bertemu",
        "dingin",
        "hangat"
    ]

    if any(keyword in q for keyword in strong_excel_keywords):
        return True

    has_analysis_intent = any(keyword in q for keyword in data_analysis_keywords)
    has_domain_context = any(keyword in q for keyword in domain_keywords)

    return has_analysis_intent and has_domain_context


def classify_task(question):
    q = question.lower().strip()

    greeting_keywords = [
        "halo",
        "hai",
        "hello",
        "selamat pagi",
        "selamat siang",
        "selamat sore",
        "selamat malam"
    ]

    text_transform_keywords = [
        "ringkas",
        "parafrase",
        "buat lebih formal",
        "buat lebih singkat",
        "judul",
        "caption",
        "ubah gaya bahasa",
        "perbaiki kalimat",
        "rapikan kalimat",
        "buatkan kalimat"
    ]

    heavy_keywords = [
        "analisis",
        "diagnosa",
        "risiko",
        "insight",
        "executive summary",
        "monthly report",
        "laporan manajemen",
        "management summary",
        "rekomendasi coaching",
        "prioritas coaching",
        "bottleneck",
        "activity tinggi",
        "cair rendah",
        "pipeline tertahan",
        "pipeline panas tinggi"
    ]

    rag_question_patterns = [
        "apa itu",
        "definisi",
        "arti dari",
        "maksud dari",
        "jelaskan stage",
        "jelaskan pipeline",
        "jelaskan activity",
        "jelaskan coaching",
        "format jawaban",
        "anti-halusinasi",
        "aturan",
        "knowledge base"
    ]

    domain_keywords = [
        "pipeline",
        "activity",
        "coaching",
        "supervisor",
        "sales",
        "tidak bertemu",
        "dingin",
        "hangat",
        "panas",
        "pemberkasan",
        "perjanjian",
        "cair"
    ]

    if any(q.startswith(keyword) for keyword in greeting_keywords):
        return "light"

    if any(keyword in q for keyword in text_transform_keywords):
        return "light"

    if any(keyword in q for keyword in heavy_keywords):
        return "heavy_rag"

    if any(pattern in q for pattern in rag_question_patterns):
        return "rag"

    if "jelaskan" in q and any(keyword in q for keyword in domain_keywords):
        return "rag"

    return "light"


def answer_streamlit_router(question, model_mode, gemini_api_key, groq_api_key, tone_mode):
    mode_map = {
        "Auto Router": "auto",
        "Excel Analyzer": "excel",
        "RAG Knowledge": "rag",
        "Groq Only": "groq",
        "Gemini Only": "gemini"
    }

    mode = mode_map.get(model_mode, "auto")

    if mode == "excel":
        if not gemini_api_key:
            return {
                "route": "excel_missing_gemini",
                "task_type": "excel_analysis",
                "answer": "Gemini API Key wajib diisi untuk menjalankan Excel Analyzer.",
                "sources": [],
                "from_cache": False
            }

        result = ask_excel_analyzer_streamlit(
            question=question,
            gemini_api_key=gemini_api_key,
            tone_mode=tone_mode,
            use_cache=True
        )

        return {
            "route": "excel_analyzer_gemini",
            "task_type": "excel_analysis",
            **result
        }

    if mode == "rag":
        if not gemini_api_key:
            return {
                "route": "rag_missing_gemini",
                "task_type": "rag",
                "answer": "Gemini API Key wajib diisi untuk menjalankan RAG Knowledge.",
                "sources": [],
                "from_cache": False
            }

        result = ask_rag_streamlit(
            question=question,
            gemini_api_key=gemini_api_key,
            tone_mode=tone_mode,
            k=4
        )

        return {
            "route": "gemini_rag",
            "task_type": "rag",
            **result
        }

    if mode == "groq":
        if not groq_api_key:
            return {
                "route": "groq_missing_key",
                "task_type": "light",
                "answer": "Groq API Key belum tersedia.",
                "sources": [],
                "from_cache": False
            }

        answer = call_groq_light_streamlit(
            question=question,
            groq_api_key=groq_api_key,
            tone_mode=tone_mode
        )

        return {
            "route": "groq_light",
            "task_type": "light",
            "answer": answer,
            "sources": [],
            "from_cache": False
        }

    if mode == "gemini":
        if not gemini_api_key:
            return {
                "route": "gemini_missing_key",
                "task_type": "gemini_direct",
                "answer": "Gemini API Key belum tersedia.",
                "sources": [],
                "from_cache": False
            }

        answer = call_gemini_direct_streamlit(
            question=question,
            gemini_api_key=gemini_api_key,
            tone_mode=tone_mode
        )

        return {
            "route": "gemini_direct",
            "task_type": "gemini_direct",
            "answer": answer,
            "sources": [],
            "from_cache": False
        }

    # AUTO ROUTER
    if should_use_excel_analyzer(question):
        if st.session_state.excel_processed and gemini_api_key:
            result = ask_excel_analyzer_streamlit(
                question=question,
                gemini_api_key=gemini_api_key,
                tone_mode=tone_mode,
                use_cache=True
            )

            return {
                "route": "excel_analyzer_gemini",
                "task_type": "excel_analysis",
                **result
            }

        if not st.session_state.excel_processed:
            return {
                "route": "excel_not_ready",
                "task_type": "excel_analysis",
                "answer": "Pertanyaan ini membutuhkan analisis Excel, tetapi Excel Analyzer belum diproses.",
                "sources": [],
                "from_cache": False
            }

        if not gemini_api_key:
            return {
                "route": "excel_missing_gemini",
                "task_type": "excel_analysis",
                "answer": "Gemini API Key wajib diisi untuk menjawab analisis Excel.",
                "sources": [],
                "from_cache": False
            }

    task_type = classify_task(question)

    if task_type in ["rag", "heavy_rag"]:
        if gemini_api_key and st.session_state.rag_ready:
            result = ask_rag_streamlit(
                question=question,
                gemini_api_key=gemini_api_key,
                tone_mode=tone_mode,
                k=4
            )

            return {
                "route": "gemini_rag",
                "task_type": task_type,
                **result
            }

        if gemini_api_key:
            answer = call_gemini_direct_streamlit(
                question=question,
                gemini_api_key=gemini_api_key,
                tone_mode=tone_mode
            )

            return {
                "route": "gemini_direct_no_rag",
                "task_type": task_type,
                "answer": answer,
                "sources": [],
                "from_cache": False
            }

        return {
            "route": "rag_or_gemini_missing_key",
            "task_type": task_type,
            "answer": "Pertanyaan ini membutuhkan Gemini/RAG, tetapi Gemini API Key belum tersedia.",
            "sources": [],
            "from_cache": False
        }

    # Light task
    if groq_api_key:
        answer = call_groq_light_streamlit(
            question=question,
            groq_api_key=groq_api_key,
            tone_mode=tone_mode
        )

        return {
            "route": "groq_light",
            "task_type": "light",
            "answer": answer,
            "sources": [],
            "from_cache": False
        }

    if gemini_api_key:
        answer = call_gemini_direct_streamlit(
            question=question,
            gemini_api_key=gemini_api_key,
            tone_mode=tone_mode
        )

        return {
            "route": "gemini_direct_no_groq",
            "task_type": "light",
            "answer": answer,
            "sources": [],
            "from_cache": False
        }

    return {
        "route": "no_api_key",
        "task_type": "none",
        "answer": "Masukkan minimal Gemini API Key atau Groq API Key untuk menjalankan chatbot.",
        "sources": [],
        "from_cache": False
    }

# =========================
# UI UTILITY FUNCTIONS
# =========================

def clear_chat_history():
    st.session_state.messages = []
    st.session_state.last_route_info = None


def clear_ai_caches():
    st.session_state.excel_answer_cache = {}
    st.session_state.router_cache = {}


def reset_excel_analyzer_state():
    st.session_state.excel_df = None
    st.session_state.excel_ready = False
    st.session_state.excel_processed = False
    st.session_state.detected_columns = {}
    st.session_state.df_pipeline = None
    st.session_state.overall_summary = None
    st.session_state.summary_sales = None
    st.session_state.summary_supervisor = None
    st.session_state.summary_region = None
    st.session_state.summary_branch = None
    st.session_state.coaching_sales = None
    st.session_state.coaching_supervisor = None
    st.session_state.rule_insights = []
    st.session_state.excel_analysis_context = None
    st.session_state.excel_answer_cache = {}


def reset_rag_state():
    st.session_state.rag_ready = False
    st.session_state.rag_vector_store = None
    st.session_state.rag_sources = []


def export_chat_history_json():
    export_data = {
        "exported_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "total_messages": len(st.session_state.messages),
        "messages": st.session_state.messages
    }

    return json.dumps(export_data, ensure_ascii=False, indent=2)


def export_chat_history_markdown():
    lines = []
    lines.append("# Sales Field Intelligence Chatbot - Chat History")
    lines.append("")
    lines.append(f"Exported at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("")

    for i, message in enumerate(st.session_state.messages, start=1):
        role = message.get("role", "unknown").upper()
        content = message.get("content", "")

        lines.append(f"## {i}. {role}")
        lines.append("")
        lines.append(content)
        lines.append("")

        route_info = message.get("route_info")

        if route_info:
            lines.append("### Route Info")
            lines.append("")
            lines.append("```json")
            lines.append(json.dumps(route_info, ensure_ascii=False, indent=2))
            lines.append("```")
            lines.append("")

        lines.append("---")
        lines.append("")

    return "\n".join(lines)


def route_badge(route):
    if not route:
        return "⚪ Unknown"

    route = str(route)

    if "excel" in route:
        return "🟢 Excel Analyzer"

    if "rag" in route:
        return "🔵 RAG Knowledge"

    if "groq" in route:
        return "🟣 Groq Light"

    if "gemini" in route:
        return "🟡 Gemini"

    if "error" in route:
        return "🔴 Error"

    return f"⚪ {route}"


def build_status_text():
    excel_status = "Ready" if st.session_state.excel_processed else "Belum Diproses"
    rag_status = "Ready" if st.session_state.rag_ready else "Belum Build"
    gemini_status = "Aktif" if st.session_state.api_status.get("gemini") else "Belum Aktif"
    groq_status = "Aktif" if st.session_state.api_status.get("groq") else "Opsional/Belum Aktif"

    text = f"""
### System Status

| Komponen | Status |
|---|---|
| Gemini | {gemini_status} |
| Groq | {groq_status} |
| Excel Analyzer | {excel_status} |
| RAG Knowledge | {rag_status} |
| Chat Messages | {len(st.session_state.messages)} |
| Excel Cache | {len(st.session_state.excel_answer_cache)} |
"""

    return text

# =========================
# SIDEBAR
# =========================

with st.sidebar:
    st.title("⚙️ Configuration")

    st.subheader("1. API Keys")

    gemini_api_key = st.text_input(
        "Gemini API Key",
        type="password",
        placeholder="Masukkan Gemini API Key"
    )

    with st.expander("Advanced - Groq API", expanded=False):
        groq_api_key = st.text_input(
            "Groq API Key",
            type="password",
            placeholder="Opsional untuk task ringan"
        )

    if gemini_api_key:
        st.session_state.api_status["gemini"] = True
        st.success("Gemini API terisi")
    else:
        st.session_state.api_status["gemini"] = False
        st.warning("Gemini API belum terisi")

    if groq_api_key:
        st.session_state.api_status["groq"] = True
        st.success("Groq API terisi")
    else:
        st.session_state.api_status["groq"] = False
        st.info("Groq API opsional")

    st.divider()

    st.subheader("2. Model Mode")

    model_mode = st.radio(
        "Pilih mode model",
        ["Auto Router", "Excel Analyzer", "RAG Knowledge", "Groq Only", "Gemini Only"],
        index=0
    )

    tone_mode = st.selectbox(
        "Gaya jawaban",
        ["Formal Operasional", "Ringkas Eksekutif", "Detail Analitis"],
        index=0
    )

    st.divider()

    st.subheader("3. Upload Excel")

    uploaded_excel = st.file_uploader(
        "Upload file Excel pipeline/activity",
        type=["xlsx", "xls"]
    )

    st.divider()

    st.subheader("4. Knowledge Base RAG")

    uploaded_kb = st.file_uploader(
        "Upload dokumen knowledge base",
        type=["pdf", "docx", "txt", "md"],
        accept_multiple_files=True
    )

    if uploaded_kb:
        st.info(f"{len(uploaded_kb)} dokumen siap diproses.")

    if st.button("Build RAG Knowledge Base"):
        if not uploaded_kb:
            st.warning("Upload minimal satu dokumen knowledge base terlebih dahulu.")
        else:
            try:
                with st.spinner("Membangun vector store RAG..."):
                    vector_store, sources, chunk_count = build_rag_vector_store_streamlit(uploaded_kb)

                    st.session_state.rag_vector_store = vector_store
                    st.session_state.rag_sources = sources
                    st.session_state.rag_ready = True

                st.success(f"RAG siap. {chunk_count} chunk dibuat dari {len(sources)} dokumen.")

            except Exception as e:
                st.session_state.rag_ready = False
                st.error(f"Gagal membangun RAG: {e}")

    st.divider()

    st.subheader("5. Utilities")

    util_col1, util_col2 = st.columns(2)

    with util_col1:
        if st.button("Clear Chat"):
            clear_chat_history()
            st.success("Chat history dihapus.")

    with util_col2:
        if st.button("Clear Cache"):
            clear_ai_caches()
            st.success("Cache dihapus.")

    util_col3, util_col4 = st.columns(2)

    with util_col3:
        if st.button("Reset Excel"):
            reset_excel_analyzer_state()
            st.success("Excel Analyzer direset.")

    with util_col4:
        if st.button("Reset RAG"):
            reset_rag_state()
            st.success("RAG direset.")

    st.download_button(
        label="Download Chat JSON",
        data=export_chat_history_json(),
        file_name="chat_history_sales_field_intelligence.json",
        mime="application/json"
    )

    st.download_button(
        label="Download Chat Markdown",
        data=export_chat_history_markdown(),
        file_name="chat_history_sales_field_intelligence.md",
        mime="text/markdown"
    )

# =========================
# MAIN HEADER
# =========================

st.title("📊 Sales Analysis Assistant")
st.caption("AI assistant untuk analisis pipeline, activity, dan rekomendasi coaching Sales/Supervisor.")

st.markdown(
    """
    <style>
    .main-title-card {
        padding: 1rem 1.2rem;
        border-radius: 14px;
        background: linear-gradient(90deg, #0F766E 0%, #134E4A 100%);
        color: white;
        margin-bottom: 1rem;
    }
    .small-muted {
        color: #6B7280;
        font-size: 0.88rem;
    }
    .route-pill {
        display: inline-block;
        padding: 0.25rem 0.6rem;
        border-radius: 999px;
        background-color: #F3F4F6;
        border: 1px solid #E5E7EB;
        font-size: 0.85rem;
        margin-top: 0.25rem;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# =========================
# STATUS CARDS
# =========================

col1, col2, col3, col4, col5 = st.columns(5)

with col1:
    st.metric(
        "Gemini",
        "Aktif" if st.session_state.api_status["gemini"] else "Belum Aktif"
    )

with col2:
    st.metric(
        "Groq",
        "Aktif" if st.session_state.api_status["groq"] else "Opsional"
    )

with col3:
    st.metric(
        "Excel",
        "Ready" if st.session_state.excel_processed else "Belum Diproses"
    )

with col4:
    st.metric(
        "RAG",
        "Ready" if st.session_state.rag_ready else "Belum Build"
    )

with col5:
    st.metric(
        "Mode",
        model_mode
    )

if st.session_state.rag_ready:
    with st.expander("📚 RAG Knowledge Base Status", expanded=False):
        st.write("Dokumen yang aktif:")
        for source in st.session_state.rag_sources:
            st.write(f"- {source}")

with st.expander("🧭 System Status Detail", expanded=False):
    st.markdown(build_status_text())

# =========================
# EXCEL UPLOAD & ANALYZER UI
# =========================

st.subheader("📄 Excel Analyzer")

if uploaded_excel is not None:
    try:
        xls = pd.ExcelFile(uploaded_excel)
        sheet_name = st.selectbox(
            "Pilih sheet Excel",
            xls.sheet_names
        )

        df = pd.read_excel(uploaded_excel, sheet_name=sheet_name)

        st.session_state.excel_df = df
        st.session_state.excel_ready = True

        st.success(f"Excel berhasil dibaca: {len(df):,} baris, {len(df.columns):,} kolom")

        with st.expander("Preview Data Excel", expanded=False):
            st.dataframe(df.head(20), use_container_width=True)

        detected = detect_columns(df)

        st.markdown("### Mapping Kolom")

        columns_options = ["-- Tidak digunakan --"] + list(df.columns)

        def default_index_for(key):
            detected_col = detected.get(key)

            if detected_col in df.columns:
                return columns_options.index(detected_col)

            return 0

        map_col1, map_col2, map_col3 = st.columns(3)

        with map_col1:
            mapped_sales = st.selectbox("Sales", columns_options, index=default_index_for("sales"))
            mapped_supervisor = st.selectbox("Supervisor", columns_options, index=default_index_for("supervisor"))
            mapped_region = st.selectbox("Region", columns_options, index=default_index_for("region"))

        with map_col2:
            mapped_area = st.selectbox("Area", columns_options, index=default_index_for("area"))
            mapped_branch = st.selectbox("Cabang/Branch", columns_options, index=default_index_for("branch"))
            mapped_stage = st.selectbox("Stage Pipeline", columns_options, index=default_index_for("stage"))

        with map_col3:
            mapped_activity = st.selectbox("Activity", columns_options, index=default_index_for("activity"))
            mapped_result = st.selectbox("Status Cair/Hasil", columns_options, index=default_index_for("result"))
            mapped_plafond = st.selectbox("Plafond", columns_options, index=default_index_for("plafond"))
            mapped_period = st.selectbox("Periode/Tanggal", columns_options, index=default_index_for("period"))
            mapped_constraint = st.selectbox("Kendala/Reason", columns_options, index=default_index_for("constraint"))

        user_mapping = {
            "sales": None if mapped_sales == "-- Tidak digunakan --" else mapped_sales,
            "supervisor": None if mapped_supervisor == "-- Tidak digunakan --" else mapped_supervisor,
            "region": None if mapped_region == "-- Tidak digunakan --" else mapped_region,
            "area": None if mapped_area == "-- Tidak digunakan --" else mapped_area,
            "branch": None if mapped_branch == "-- Tidak digunakan --" else mapped_branch,
            "stage": None if mapped_stage == "-- Tidak digunakan --" else mapped_stage,
            "activity": None if mapped_activity == "-- Tidak digunakan --" else mapped_activity,
            "result": None if mapped_result == "-- Tidak digunakan --" else mapped_result,
            "plafond": None if mapped_plafond == "-- Tidak digunakan --" else mapped_plafond,
            "period": None if mapped_period == "-- Tidak digunakan --" else mapped_period,
            "constraint": None if mapped_constraint == "-- Tidak digunakan --" else mapped_constraint,
        }

        with st.expander("Detected Column Mapping", expanded=False):
            st.json(user_mapping)

        if st.button("🚀 Proses Excel Analyzer", type="primary"):
            if user_mapping["stage"] is None:
                st.error("Kolom Stage Pipeline wajib dipilih.")
            else:
                result = run_excel_analyzer(df, user_mapping)

                st.session_state.detected_columns = user_mapping
                st.session_state.df_pipeline = result["df_pipeline"]
                st.session_state.overall_summary = result["overall_summary"]
                st.session_state.summary_sales = result["summary_sales"]
                st.session_state.summary_supervisor = result["summary_supervisor"]
                st.session_state.summary_region = result["summary_region"]
                st.session_state.summary_branch = result["summary_branch"]
                st.session_state.rule_insights = result["rule_insights"]
                st.session_state.coaching_sales = result["coaching_sales"]
                st.session_state.coaching_supervisor = result["coaching_supervisor"]
                st.session_state.excel_analysis_context = result["excel_analysis_context"]
                st.session_state.excel_processed = True

                st.success("Excel Analyzer berhasil diproses.")

    except Exception as e:
        st.session_state.excel_ready = False
        st.session_state.excel_processed = False
        st.error(f"Gagal memproses Excel: {e}")

else:
    st.info("Upload file Excel untuk mengaktifkan Excel Analyzer.")

# =========================
# EXCEL ANALYZER RESULT
# =========================

if st.session_state.excel_processed:
    overall = st.session_state.overall_summary

    st.markdown("## 📊 Quick Summary")

    q1, q2, q3, q4, q5 = st.columns(5)

    with q1:
        st.metric("Total Data", f"{overall['total_rows']:,}")

    with q2:
        st.metric("Total Activity", f"{overall['total_activity']:,.0f}")

    with q3:
        st.metric("Jumlah Cair", f"{overall['total_cair']:,}")

    with q4:
        st.metric("Pipeline Matang", f"{overall['mature_count']:,}")

    with q5:
        st.metric("Pipeline Awal", f"{overall['early_count']:,}")

    st.markdown("### Rule-Based Insight")

    for insight in st.session_state.rule_insights:
        st.info(insight)

    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "Stage Summary",
        "Region",
        "Sales Coaching",
        "Supervisor Coaching",
        "AI Context Preview"
    ])

    with tab1:
        st.dataframe(overall["stage_summary"], use_container_width=True)

    with tab2:
        st.dataframe(st.session_state.summary_region.head(20), use_container_width=True)

    with tab3:
        st.dataframe(st.session_state.coaching_sales, use_container_width=True)

    with tab4:
        st.dataframe(st.session_state.coaching_supervisor, use_container_width=True)

    with tab5:
        st.text_area(
            "Excel Analysis Context",
            st.session_state.excel_analysis_context,
            height=400
        )

# =========================
# PROMPT LIBRARY
# =========================

st.divider()
st.subheader("🧪 Prompt Testing Library")

with st.expander("Klik contoh pertanyaan untuk testing router", expanded=False):
    p1, p2, p3 = st.columns(3)

    with p1:
        if st.button("Ringkasan Excel"):
            st.session_state.pending_prompt = "Berdasarkan data Excel, buatkan ringkasan pipeline dan activity."

        if st.button("Prioritas Coaching Sales"):
            st.session_state.pending_prompt = "Sales mana yang perlu menjadi prioritas coaching berdasarkan data Excel?"

    with p2:
        if st.button("Region Perlu Perhatian"):
            st.session_state.pending_prompt = "Region mana yang perlu perhatian dari data ini?"

        if st.button("Definisi Panas"):
            st.session_state.pending_prompt = "Apa itu Panas dalam pipeline?"

    with p3:
        if st.button("Sapaan / Light Task"):
            st.session_state.pending_prompt = "Halo, jelaskan singkat chatbot ini bisa bantu apa?"

        if st.button("Parafrase"):
            st.session_state.pending_prompt = "Parafrase kalimat ini agar lebih formal: Sales harus segera follow up nasabah yang belum cair."

# =========================
# CHAT AREA
# =========================

st.divider()
st.subheader("💬 Chat")

for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

        route_info = message.get("route_info")

        if route_info:
            badge = route_badge(route_info.get("route"))
            st.markdown(
                f"<span class='route-pill'>{badge}</span>",
                unsafe_allow_html=True
            )

            with st.expander("Route Info", expanded=False):
                st.json(route_info)

pending_prompt = st.session_state.get("pending_prompt")
chat_prompt = st.chat_input("Tanyakan analisis pipeline, activity, coaching, atau knowledge base...")

user_prompt = pending_prompt or chat_prompt

if pending_prompt:
    st.session_state.pending_prompt = None

if user_prompt:
    st.session_state.messages.append(
        {
            "role": "user",
            "content": user_prompt
        }
    )

    with st.chat_message("user"):
        st.markdown(user_prompt)

    with st.chat_message("assistant"):
        with st.spinner("Menganalisis..."):
            try:
                result = answer_streamlit_router(
                    question=user_prompt,
                    model_mode=model_mode,
                    gemini_api_key=gemini_api_key,
                    groq_api_key=groq_api_key,
                    tone_mode=tone_mode
                )

                response = result.get("answer", "")
                route_info = {
                    "route": result.get("route"),
                    "route_badge": route_badge(result.get("route")),
                    "task_type": result.get("task_type"),
                    "from_cache": result.get("from_cache"),
                    "sources": result.get("sources", [])
                }

                st.markdown(response)

                st.markdown(
                    f"<span class='route-pill'>{route_info['route_badge']}</span>",
                    unsafe_allow_html=True
                )

                with st.expander("Route Info", expanded=False):
                    st.json(route_info)

                sources = result.get("sources", [])

                if sources:
                    with st.expander("Sumber Referensi", expanded=False):
                        for source in sources:
                            st.write(f"- {source}")

            except Exception as e:
                response = f"Terjadi error saat memproses pertanyaan: {e}"
                route_info = {
                    "route": "error",
                    "route_badge": "🔴 Error",
                    "task_type": "error",
                    "from_cache": False,
                    "sources": []
                }

                st.error(response)

    st.session_state.messages.append(
        {
            "role": "assistant",
            "content": response,
            "route_info": route_info
        }
    )

    st.session_state.last_route_info = route_info
